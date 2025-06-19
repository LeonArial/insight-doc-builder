import os
import re
import copy
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def replace_text_in_element(element, replacements):
    """在一个元素（如文档、页眉、单元格）内递归替换文本。"""
    for para in element.paragraphs:
        # 为避免破坏样式，我们逐个 run 进行替换
        for key, value in replacements.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))

    for table in element.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_element(cell, replacements)

def replace_general_placeholders(doc, replacements):
    """替换整个文档中的通用文本占位符（正文、表格、页眉、页脚）。"""
    replace_text_in_element(doc, replacements)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                replace_text_in_element(header, replacements)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                replace_text_in_element(footer, replacements)

def clear_paragraph(p):
    """清空一个段落的所有内容。"""
    for run in p.runs:
        run.text = ''

def clear_cell_content(cell):
    """清空一个单元格的所有内容。"""
    for p in cell.paragraphs:
        clear_paragraph(p)

def populate_vulnerabilities(doc, vulnerabilities, system_name):
    """
    Dynamically creates vulnerability entries based on a single template
    in the document and populates them.
    """
    # 1. Populate Summary Table
    summary_table = next((tbl for tbl in doc.tables if len(tbl.rows) > 0 and len(tbl.columns) > 4 and '系统名称' in tbl.cell(0, 1).text and '漏洞名称' in tbl.cell(0, 2).text), None)
    if summary_table:
        for i in range(len(summary_table.rows) - 1, 0, -1):
            row = summary_table.rows[i]
            row._element.getparent().remove(row._element)
        for i, vuln in enumerate(vulnerabilities, 1):
            cells = summary_table.add_row().cells
            cells[0].text = str(i)
            cells[1].text = system_name
            cells[2].text = vuln['name']
            cells[3].text = vuln['risk']
            cells[4].text = vuln['description']
    else:
        print("警告: 未找到漏洞概要表。")

    # 2. Find the single vulnerability template (paragraph + table)
    p_template = next((p for p in doc.paragraphs if '【#VULNERABILITIES_SECTION#】' in p.text), None)
    t_template = next((tbl for tbl in doc.tables if len(tbl.rows) > 0 and '风险描述' in tbl.cell(0, 0).text), None)

    if not p_template or not t_template:
        if not vulnerabilities:
            if p_template: clear_paragraph(p_template)
            if t_template:
                for row in t_template.rows:
                    for cell in row.cells:
                        clear_cell_content(cell)
        print("警告: 未能在模板中找到漏洞详情占位符段落或表格。")
        return

    # 3. If there are more vulnerabilities than the template slot, create new slots by cloning.
    if len(vulnerabilities) > 1:
        p_template_xml = p_template._p
        t_template_xml = t_template._tbl
        last_element = t_template_xml
        for i in range(len(vulnerabilities) - 1):
            new_p = copy.deepcopy(p_template_xml)
            new_t = copy.deepcopy(t_template_xml)
            last_element.addnext(new_p)
            new_p.addnext(new_t)
            last_element = new_t

    # 4. Populate all vulnerability slots (original template + new ones)
    all_title_paragraphs = [p for p in doc.paragraphs if '【#VULNERABILITIES_SECTION#】' in p.text]
    all_detail_tables = [tbl for tbl in doc.tables if len(tbl.rows) > 0 and '风险描述' in tbl.cell(0, 0).text]

    if not vulnerabilities:
        if all_title_paragraphs:
            clear_paragraph(all_title_paragraphs[0])
        if all_detail_tables:
            for row in all_detail_tables[0].rows:
                for cell in row.cells:
                    clear_cell_content(cell)
        return

    num_slots_to_fill = min(len(vulnerabilities), len(all_title_paragraphs), len(all_detail_tables))
    for i in range(num_slots_to_fill):
        vuln = vulnerabilities[i]
        p = all_title_paragraphs[i]
        table = all_detail_tables[i]

        # a. Populate title paragraph
        prefix = f"5.1.{i+1}"
        clear_paragraph(p)
        p.text = f"{prefix} 【{vuln['risk']}】{vuln['name']}"

        # b. Populate detail table
        clear_cell_content(table.cell(0, 1))
        table.cell(0, 1).text = vuln['description']
        
        cell = table.cell(1, 1)
        clear_cell_content(cell)
        # Paragraph for text (left-aligned by default)
        p_text = cell.paragraphs[0]
        p_text.add_run(vuln['process']['text'])

        # Add a new, centered paragraph for the image
        image_path = vuln['process']['image_path']
        if image_path and os.path.exists(image_path):
            try:
                # Add a new paragraph specifically for the image
                p_image = cell.add_paragraph()
                run = p_image.add_run()
                run.add_picture(image_path, width=Inches(5.0))
                p_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                # If image fails, add error message in a new paragraph
                cell.add_paragraph(f'\n[图片添加失败: {e}]')
        
        clear_cell_content(table.cell(2, 1))
        table.cell(2, 1).text = vuln['advice']

def main():
    """主函数：生成报告"""
    template_path = '渗透报告模板.docx'
    output_path = f'渗透测试报告_{datetime.datetime.now().strftime("%Y%m%d%H%M%S")}.docx'

    replacements = {
        '#SYSTEM#': 'XX管理信息系统',
        '#DATA#': datetime.datetime.now().strftime("%Y/%m/%d"),
        '#IP#': '192.168.1.100',
        '#DATARANGE#': '2025/06/17-2025/06/19',
        '#DJNAME#': '张三',
        '#DJPHONE#': '13800138000',
        '#CSNAME#': '李四',
        '#CSPHONE#': '13900139000',
        '#CSACCOUNT#': 'admin',
    }

    vulnerabilities = [
        {
            'risk': '高危',
            'name': 'SQL注入漏洞',
            'description': '在用户登录处发现SQL注入漏洞，攻击者可能利用该漏洞绕过认证，获取数据库敏感信息。',
            'process': {
                'text': '1. 访问登录页面。\n2. 在用户名输入框中输入恶意构造的SQL查询语句 `\' OR \'1\'=\'1`。\n3. 系统成功登录，暴露了SQL注入风险，如下图所示：',
                'image_path': 'poc_sql_injection.png'
            },
            'advice': '建议使用参数化查询或ORM框架来处理数据库操作，并对所有用户输入进行严格的验证和过滤。'
        },
        {
            'risk': '中危',
            'name': '跨站脚本攻击 (XSS)',
            'description': '在搜索功能处发现存储型XSS漏洞，攻击者可注入恶意脚本，窃取其他用户的Cookie或执行恶意操作。',
            'process': {
                'text': '1. 访问搜索页面。\n2. 在搜索框中输入恶意脚本 `<script>alert("XSS")</script>` 并提交。\n3. 当其他用户访问搜索结果页面时，恶意脚本被执行，弹出警告框，如下图所示：',
                'image_path': 'poc_xss.png'
            },
            'advice': '建议对所有用户输入及输出到页面的内容进行严格的HTML实体编码，特别是对 `<, >, ", \'` 等特殊字符。'
        }
    ]

    if not os.path.exists(template_path):
        print(f"错误：模板文件 '{template_path}' 不存在。")
        return
    
    doc = Document(template_path)

    # 步骤1: 替换通用占位符
    replace_general_placeholders(doc, replacements)

    # 步骤2: 填充漏洞相关内容
    populate_vulnerabilities(doc, vulnerabilities, replacements['#SYSTEM#'])

    try:
        doc.save(output_path)
        print(f"报告生成成功！已保存为：{output_path}")
    except Exception as e:
        print(f"保存文件时出错: {e}")

if __name__ == '__main__':
    main()
