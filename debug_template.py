import datetime
from docx import Document

def analyze_document_to_md(doc_path, output_path):
    """
    分析Word文档的结构和内容，并将结果保存为Markdown文件。
    """
    md_lines = []
    try:
        doc = Document(doc_path)
        md_lines.append(f"# 文档分析报告: `{doc_path}`")
        md_lines.append(f"> 生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        md_lines.append("\n---\n")

        md_lines.append("## 正文内容 (Paragraphs)")
        for i, para in enumerate(doc.paragraphs):
            md_lines.append(f"### 段落 {i+1}")
            # 使用代码块来显示原始文本，防止Markdown解析
            md_lines.append(f"```text\n{para.text}\n```")
            
            # 检查一个段落是否被分割成多个 'run'，这对于调试占位符很重要
            if len(para.runs) > 1 and any('#' in run.text for run in para.runs):
                 run_texts = ' | '.join([f"`{run.text}`" for run in para.runs])
                 md_lines.append(f"> **[!] 注意：** 这个包含占位符的段落被分割成了 {len(para.runs)} 个部分: {run_texts}")
            md_lines.append("\n")

        md_lines.append("\n---\n")
        md_lines.append("## 表格内容 (Tables)")
        for i, table in enumerate(doc.tables):
            md_lines.append(f"\n### 表格 {i+1}")
            for r, row in enumerate(table.rows):
                md_lines.append(f"**行 {r+1}:**")
                row_contents = []
                for c, cell in enumerate(row.cells):
                    # 将单元格内容中的换行符替换为 <br> 以便在Markdown中显示
                    cell_text = cell.text.replace('\n', '<br>')
                    row_contents.append(f"_{c+1}_: `{cell_text}`")
                md_lines.append(' | '.join(row_contents))
            md_lines.append("\n")

        md_lines.append("\n---\n")
        md_lines.append("分析完成。")
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(md_lines))
        print(f"分析报告已成功生成: {output_path}")

    except Exception as e:
        print(f"打开或分析文档时出错: {e}")

if __name__ == '__main__':
    template_path = '渗透报告模板.docx'
    output_md_path = 'template_analysis.md'
    analyze_document_to_md(template_path, output_md_path)
